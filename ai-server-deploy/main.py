import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from flask import Flask, request, jsonify
from flask_cors import CORS
from PIL import Image
from pdf2image import convert_from_bytes
import os
import json
import io
import cv2
import numpy as np
import time
import shutil
import sys
import uuid
import threading
import concurrent.futures
import platform
import google.api_core.exceptions
import re
from queue import Queue
from dotenv import load_dotenv

load_dotenv()

# Reconfigure stdout for utf-8 logs in Cloud Run
sys.stdout.reconfigure(encoding='utf-8')

# ==========================================
# API KEY
# ==========================================
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError('GEMINI_API_KEY not set in environment variables')

genai.configure(api_key=GEMINI_API_KEY)
for m in genai.list_models():
    print(m.name)

# ==========================================
# SELECT MODELS
# ==========================================
transcription_model = genai.GenerativeModel('gemini-2.5-pro')
grading_model = genai.GenerativeModel('gemini-2.5-pro')
print(f"🎯 TRANSCRIPTION MODEL:", transcription_model.model_name)
print(f"🎯 GRADING MODEL:", grading_model.model_name)

app = Flask(__name__)
CORS(app)

# ==========================================
# OUTPUT FOLDERS
# ==========================================
TRANSCRIPT_DIR = "processed_images/transcripts"
GRADE_LOG_DIR = "processed_images/grade_logs"

for d in [TRANSCRIPT_DIR, GRADE_LOG_DIR]:
    if not os.path.exists(d):
        os.makedirs(d)

# ==========================================
# SAFETY SETTINGS & GENERATION CONFIG
# ==========================================
SAFETY_SETTINGS = {
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

GENERATION_CONFIG = genai.GenerationConfig(
    temperature=0,
    top_p=1,
    top_k=1,
)


# ==========================================
# PROMPTS (defined once, reused everywhere)
# ==========================================
TRANSCRIPTION_PROMPT = """
You are a document digitization tool.
Your only task is to read and copy handwritten text from the provided image exactly as written.

RULES:
- Transcribe every word exactly as written.
- Preserve all spelling errors, grammar mistakes, and punctuation.
- Do NOT correct anything.
- Do NOT grade or evaluate anything.
- Mark unclear words as [unclear: best_guess]
- Mark fully unreadable lines as [unclear]
- Preserve question numbers exactly as written (e.g. 46. 47. 48.)
- Transcribe ALL numbered questions present in the image.

MULTIPLE CHOICE DETECTION RULES:
- If a letter option (A, B, C, D) is CIRCLED, UNDERLINED, BOXED, or visually marked
  next to a question number — treat that as the student's answer.
- Write it as: "1. B" (question number + selected letter).
- Do NOT write "BLANK" if you can see a circled or marked letter.
- A faint circle, partial circle, or checkmark next to a letter still counts as selected.
- If truly nothing is marked, then write BLANK.

CONFIDENCE RULES:
- Start at 100
- Deduct 20 if handwriting is mostly unclear
- Deduct 40 if handwriting is very unclear
- Deduct 15 per [unclear] or [unclear: best_guess] token
- Minimum: 0

Return ONLY this JSON, no extra text:
{
  "transcribed_text": "full transcription preserving all question numbers and answers",
  "legibility": "CLEAR / MOSTLY_CLEAR / UNCLEAR",
  "confidence_score": <0-100>
}
""".strip()

GRADING_PROMPT = """
Target: Deterministic Grading Engine
Input Context: {context_block}
Student Transcription: {transcribed_text}

[CRITICAL RULE — PAGE-SCOPED GRADING]
You are grading ONE PAGE of a multi-page exam.
- ONLY grade questions that are EXPLICITLY PRESENT in the Student Transcription above.
- Do NOT invent, assume, or carry over questions from the answer key that are not visible in this transcription.
- If a question number does not appear in the transcription, SKIP IT entirely — do not mark it as BLANK.
- If a section header appears in the answer key but no questions from that section appear in the transcription, OMIT that section from your log entirely.

[GRADING RULES]
1. Detect Sections: Separate ESSAY (Explain/Discuss) from OBJECTIVE (MCQ/TF/Matching).
2. Objective: 1pt if correct (case-insensitive, substring allowed). 0pt if wrong.
3. Essay: Score = floor((Present/Total) * Max). Apply rubric minimum for any attempt.
4. Logic: Never invent quotes. If unsure, score lower. Return ONLY JSON.

OBJECTIVE SCORING RULES (MCQ / TRUE-FALSE / MATCHING):
- Compare each student answer to the answer key exactly.
- Correct = 1 point. Wrong or blank = 0 points.
- No partial credit. No weighting.
- Raw score = total correct answers only.
- SKIP any question whose correct answer is "N/A" — these are essay questions.

MATCHING TYPE RULES:
- Each correctly matched pair = 1 point.
- Treat each numbered item in the matching section as a separate answer.
- Compare the student's chosen match to the answer key for that item.
- If the student writes the wrong letter/word or leaves it blank = 0 points.
- Do NOT give partial credit for "close" matches.
- Count ALL matching items present in the answer key.

MATCHING LENIENCY RULES (apply BEFORE marking wrong):
- Ignore minor spelling errors (e.g. "Principalia" vs "Principalias" = CORRECT).
- If the student answer is a SUBSTRING of the correct answer = CORRECT.
- If the correct answer is a SUBSTRING of the student answer = CORRECT.
- If the student writes multiple answers separated by "/" or "," and ANY ONE matches = CORRECT.
- Case-insensitive. Accents and punctuation differences are ignored.
- Only mark WRONG if there is NO reasonable match after applying all leniency rules.

OBJECTIVE SCORE LOG FORMAT:
Group questions by their section (MULTIPLE CHOICE, TRUE OR FALSE, IDENTIFICATION, MATCHING TYPE, etc.)
Print the section name as a header before its questions.
After all questions in a section, print the section subtotal.
Separate sections with a divider line.

Example:
[MULTIPLE CHOICE]
Question: 1
Student Answer: B
Correct Answer: B
Points: 1
---
Question: 2
Student Answer: A
Correct Answer: C
Points: 0
---
SECTION SCORE: 1 / 2
=====================================

[TRUE OR FALSE]
Question: 11
Student Answer: TRUE
Correct Answer: True
Points: 1
---
SECTION SCORE: 1 / 1
=====================================

[OUTPUT LOG FORMATS]
- ESSAY_LOG:
Question: [number]
Rubric Requirements: [list]
Present (P/R):
- [req] → "[exact quote]"
Absent:
- [req] → no direct quote found
Rubric Max Score: [max]
Final Score: floor((P/R) * max) = [result]
Reason: [1-2 sentences]

After all sections, write:
TOTAL OBJECTIVE SCORE: {{exact number}}
=====================================
TOTAL ESSAY SCORE: {{exact number}}
=====================================
FINAL TOTAL SCORE: {{exact number}}

FAIL-SAFE RULES:
- Missing rubric → score 0.
- Blank answer → score 0.
- Never fabricate quotes or requirements.
- Always use floor().
- Never double-count a question in both objective and essay logs.
- Never re-evaluate or override the floor() calculation.
- Final Score is always floor((P/R) * max). No exceptions.
- Do not add "Final Score Re-evaluation" sectio

Return ONLY this JSON, no extra text:
{{
  "grading_type": "STRICT_MATCH or RUBRIC_BASED",
  "objective_score_log": "Plain text. Use \\n for line breaks.",
  "essay_score_log": "Plain text. Use \\n for line breaks.",
  "confidence_score": <0-100>,
  "objective_total": <numeric — TOTAL OBJECTIVE SCORE>,
  "essay_total": <numeric — TOTAL ESSAY SCORE>,
  "score": <number>,
  "feedback": "brief feedback."
}}
""".strip()

# ==========================================
# IMAGE HELPERS
# ==========================================
def enhance_image(pil_img):
    """Denoise and sharpen a PIL image using OpenCV."""
    try:
        img = np.array(pil_img)
        if len(img.shape) == 3:
            img = img[:, :, ::-1].copy()

        denoised = cv2.fastNlMeansDenoisingColored(img, None, 10, 10, 7, 21)
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        sharpened = cv2.filter2D(denoised, -1, kernel)

        return Image.fromarray(cv2.cvtColor(sharpened, cv2.COLOR_BGR2RGB))
    except Exception as e:
        print(f"⚠️ OpenCV Error: {e}")
        return pil_img


def extract_images_to_text(images: list, label: str = "Reference") -> str:
    """
    Use Flash to extract text from reference/rubric images once,
    so they don't need to be re-sent as images to every grading call.
    """
    if not images:
        return ""

    extracted_parts = []
    for i, img in enumerate(images):
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=90)

        prompt = """Extract ALL text from this image exactly as written.
Preserve formatting, numbering, and structure.
This may be an answer key, rubric, or reference material.
Return ONLY the extracted text, no commentary."""

        try:
          response = transcription_model.generate_content(
            [prompt, {"mime_type": "image/jpeg", "data": buf.getvalue()}],
            generation_config=GENERATION_CONFIG,
            safety_settings=SAFETY_SETTINGS,
          )
          text = response.text.strip()
          if text:
              extracted_parts.append(f"[{label} Image {i+1}]\n{text}")
              print(f"✅ Extracted {label} image {i+1}: {len(text)} chars")
        except Exception as e:
            print(f"⚠️ Could not extract {label} image {i+1}: {e}")

    return "\n\n".join(extracted_parts)


def read_file_as_images(file_bytes, label="exam"):
    """Convert uploaded file bytes to list of enhanced PIL images."""
    raw_images = []
    try:
        raw_images = convert_from_bytes(file_bytes)
        print(f"📄 PDF Detected ({label}): {len(raw_images)} pages")
    except Exception:
        try:
            img = Image.open(io.BytesIO(file_bytes))
            print(f"🖼️ Image Detected ({label}): {img.mode}")
            raw_images = [img]
        except Exception:
            raise ValueError(f"Invalid file format for {label}. Supported: PDF, JPG, PNG.")

    return [enhance_image(img) for img in raw_images]


# ==========================================
# FILE HELPERS
# ==========================================
def extract_docx_text(file_bytes):
    """Extract all text from a DOCX file including tables."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    table_texts.append(cell_text)

    return "\n".join(paragraphs + table_texts)


def fetch_and_parse_url(url):
    import requests as req
    from urllib.parse import urlparse, unquote

    response = req.get(url, timeout=15)

    if response.status_code != 200:
        print(f"⚠️ URL returned HTTP {response.status_code}: {url[:80]}")
        return [], ""

    file_bytes = response.content
    content_type = response.headers.get("Content-Type", "")
    parsed = urlparse(url)                                     
    key_name = unquote(parsed.path.split("/")[-1]).lower()     
    print(f"📎 Processing URL file: {key_name} | {content_type} | {len(file_bytes)} bytes")

    if "text/html" in content_type or file_bytes[:5] in (b"<!DOC", b"<html"):
        print(f"⚠️ URL returned HTML — likely an unauthenticated storage path")
        return [], ""

    if key_name.endswith('.docx') or key_name.endswith('.doc'):
        if file_bytes[:4] != b'PK\x03\x04':
            print(f"⚠️ Not a valid DOCX (got: {file_bytes[:8]})")
            return [], ""
        try:
            text = extract_docx_text(file_bytes)
            print(f"📄 DOCX extracted: {len(text)} chars")
            return [], text
        except Exception as e:
            print(f"⚠️ DOCX extraction failed: {e}")
            return [], ""

    if key_name.endswith('.pdf'):
        print(f"🔍 PDF bytes: {file_bytes[:8]} | size: {len(file_bytes)}")
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(
                page.extract_text() for page in reader.pages 
                if page.extract_text()
            )
            if text.strip():
                print(f"📄 PDF text extracted: {len(text)} chars")
                return [], text  # ✅ no Gemini call needed
            else:
                print(f"⚠️ PDF has no extractable text (likely scanned image-only PDF)")
                return [], ""
        except Exception as e:
            print(f"⚠️ PDF extraction failed: {e}")
            return [], ""

    print(f"⚠️ Unsupported URL file format: {key_name}")
    return [], ""


# ==========================================
# AI HELPERS
# ==========================================
def call_gemini(ai_inputs, model):
    response = model.generate_content(
        ai_inputs,
        generation_config=GENERATION_CONFIG,
        safety_settings=SAFETY_SETTINGS,
    )
    
    usage = getattr(response, 'usage_metadata', None)
    if usage:
        print(f"🔢 Tokens — input: {usage.prompt_token_count}, output: {usage.candidates_token_count}, total: {usage.total_token_count}")


    if not response.candidates:
        raise ValueError("AI Safety Block — no candidates returned.")

    candidate = response.candidates[0]
    finish_reason = candidate.finish_reason

    if finish_reason == 3:
        blocked = [
            f"{r.category.name}: {r.probability.name}"
            for r in candidate.safety_ratings
            if r.probability.name not in ("NEGLIGIBLE", "LOW")
        ]
        raise ValueError(f"Safety block. Ratings: {blocked}")

    if not response.parts:
        raise ValueError(f"No response parts. Finish reason: {finish_reason}")

    raw_text = response.text
    start = raw_text.find('{')
    end = raw_text.rfind('}') + 1

    if start == -1:
        raise ValueError(f"AI returned invalid JSON: {raw_text[:200]}")
    
    json_str = raw_text[start:end]

      # First attempt: parse as-is
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        pass

    # Second attempt: fix unescaped quotes inside string values
    try:
        import re
        # Remove control characters that break JSON
        json_str = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', json_str)
        # ✅ Replace unescaped newlines without lookbehind
        json_str = json_str.replace('\r\n', '\\n').replace('\r', '\\n').replace('\n', '\\n')
        return json.loads(json_str)
    except json.JSONDecodeError:
        pass

    # Third attempt: extract just the string values using a more lenient approach
    try:
        import re
        result = {}
        for key in ["transcribed_text", "legibility", "confidence_score",
                    "grading_type", "objective_score_log", "essay_score_log",
                    "objective_total", "essay_total", "score", "feedback"]:
            pattern = rf'"{key}"\s*:\s*(".*?"|[\d.]+|true|false|null)'
            match = re.search(pattern, json_str, re.DOTALL)
            if match:
                try:
                    result[key] = json.loads(match.group(1))
                except Exception:
                    result[key] = match.group(1).strip('"')
        if result:
            return result
    except Exception:
        pass


    raise ValueError(f"AI returned unparseable JSON: {json_str[:300]}")


def save_transcript_log(data):
    """Save a transcript verification report to disk."""
    try:
        timestamp = int(time.time())
        filename = f"{TRANSCRIPT_DIR}/exam_{timestamp}_verification.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"--- VERIFICATION REPORT ({timestamp}) ---\n")
            f.write(f"CONFIDENCE SCORE: {data.get('confidence_score', 'N/A')}%\n")
            f.write("-" * 30 + "\n")
            f.write("TRANSCRIPTION:\n")
            f.write(data.get("transcribed_text", ""))
        print(f"💾 Saved: {filename}")
    except Exception as e:
        print(f"⚠️ Could not save transcript: {e}")
        
def save_grade_log(data):
    """Save objective and essay score logs to disk."""
    try:
        timestamp = int(time.time())
        filename = f"{GRADE_LOG_DIR}/grade_{timestamp}.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"--- GRADE LOG ({timestamp}) ---\n")
            f.write(f"SCORE: {data.get('score', 'N/A')}\n")
            f.write(f"GRADING TYPE: {data.get('grading_type', 'N/A')}\n")
            f.write(f"CONFIDENCE: {data.get('confidence_score', 'N/A')}%\n")
            f.write("-" * 30 + "\n")
            f.write("OBJECTIVE SCORE LOG:\n")
            f.write(data.get("objective_score_log", "N/A"))
            f.write("\n" + "-" * 30 + "\n")
            f.write("ESSAY SCORE LOG:\n")
            f.write(data.get("essay_score_log", "N/A"))
            f.write("\n" + "-" * 30 + "\n")
            f.write("FEEDBACK:\n")
            f.write(data.get("feedback", "N/A"))
        print(f"💾 Saved grade log: {filename}")
    except Exception as e:
        print(f"⚠️ Could not save grade log: {e}")

def create_context_cache(context_block: str) -> str | None:
    """
    Upload the context_block to Gemini's cache.
    Returns cache name (string) or None if caching failed/not worth it.
    """
    # Only cache if context is large enough to be worth it
    estimated_tokens = len(context_block) // 4
    if estimated_tokens < 1024:
        print(f"⚠️ Context too small to cache ({estimated_tokens} est. tokens), skipping.")
        return None

    try:
        import datetime
        cache = genai.caching.CachedContent.create(
            model='models/gemini-2.5-pro',
            contents=[context_block],
            ttl=datetime.timedelta(minutes=60),  # Cache lives 60 mins
        )
        print(f"✅ Context cached: {cache.name} | tokens: {cache.usage_metadata.total_token_count}")
        return cache.name
    except Exception as e:
        print(f"⚠️ Context caching failed: {e}. Falling back to inline context.")
        return None


def delete_context_cache(cache_name: str):
    """Clean up the cache after grading is done."""
    try:
        cache = genai.caching.CachedContent.get(cache_name)
        cache.delete()
        print(f"🗑️ Cache deleted: {cache_name}")
    except Exception as e:
        print(f"⚠️ Could not delete cache {cache_name}: {e}")


# ==========================================
# TRANSCRIPTION HELPERS (per-page + merge)
# ==========================================
def transcribe_single_page(img, page_num):
    """Transcribe one page image and return result tagged with page number."""
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=90)
    data = call_gemini(
        [TRANSCRIPTION_PROMPT, {"mime_type": "image/jpeg", "data": buf.getvalue()}],
        transcription_model,
    )
    
    data.setdefault("transcribed_text", "")
    data.setdefault("legibility", "UNCLEAR")
    data.setdefault("confidence_score", 0)
    
    data["page"] = page_num
    print(f"✅ Page {page_num} transcribed: {len(data.get('transcribed_text', ''))} chars")
    return data


def merge_transcriptions(page_results):
    """Merge per-page transcription results into one combined result."""
    page_results.sort(key=lambda x: x["page"])

    combined_text = "\n\n".join(
        f"[PAGE {r['page']}]\n{r['transcribed_text']}"
        for r in page_results
    )
    legibilities = [r["legibility"] for r in page_results]
    worst = (
        "UNCLEAR" if "UNCLEAR" in legibilities
        else "MOSTLY_CLEAR" if "MOSTLY_CLEAR" in legibilities
        else "CLEAR"
    )
    avg_confidence = sum(r["confidence_score"] for r in page_results) // len(page_results)

    return {
        "transcribed_text": combined_text,
        "legibility": worst,
        "confidence_score": avg_confidence,
        "pages": len(page_results)
    }

def grade_single_page(transcribed_text, context_block, page_num, cache_name: str | None = None):
    
    if page_num == 1:
        identifiers = ["gem 14", "gem-14", "life and works", "jose rizal", "rizal"]
        text_lower = transcribed_text.lower()
        if not any(keyword in text_lower for keyword in identifiers):
            raise ValueError("Wrong subject — this grader only accepts GEM 14 exams.")
        
    prompt = GRADING_PROMPT.format(
        context_block=context_block,
        transcribed_text=transcribed_text
    )

    if cache_name:
        # Use cached content model instead of base model
        cached_model = genai.GenerativeModel.from_cached_content(
            cached_content=cache_name
        )
        data = call_gemini([prompt], cached_model)
    else:
        data = call_gemini([prompt], grading_model)

    data["page"] = page_num
    print(f"✅ Page {page_num} graded: score={data.get('score')}")
    return data

def extract_essay_total_from_log(essay_log: str) -> int:
    """
    Parse 'Final Score: N' or 'Final Score: floor(...) = N' lines from essay log.
    """
    # Match the LAST number on any 'Final Score:' line — handles both formats
    scores = re.findall(r'Final Score:[^\n]*?(\d+)\s*$', essay_log, re.MULTILINE)
    return sum(int(s) for s in scores)


def merge_grades(page_grades):
    page_grades.sort(key=lambda x: x["page"])

    combined_objective_log = "\n\n".join(
        g.get('objective_score_log', '')
        for g in page_grades
        if g.get('objective_score_log', '').strip()
    )
    combined_essay_log = "\n\n".join(
        g.get('essay_score_log', '')
        for g in page_grades
        if g.get('essay_score_log', '').strip()
    )

    objective_total = sum(g.get("objective_total", 0) or 0 for g in page_grades)

    essay_total_from_json = sum(g.get("essay_total", 0) or 0 for g in page_grades)
    essay_total_from_log = extract_essay_total_from_log(combined_essay_log)
    essay_total = max(essay_total_from_json, essay_total_from_log)

    if essay_total_from_json != essay_total_from_log:
        print(f"⚠️ Essay total mismatch — JSON: {essay_total_from_json}, Log: {essay_total_from_log}, Using: {essay_total}")

    avg_confidence = sum(g.get("confidence_score", 0) or 0 for g in page_grades) // len(page_grades)
    feedbacks = [g.get("feedback", "") for g in page_grades if g.get("feedback")]

    #Append totals to the log strings so frontend can find them
    combined_objective_log += f"\n\n=====================================\nTOTAL OBJECTIVE SCORE: {objective_total}"
    combined_essay_log += f"\n\n=====================================\nTOTAL ESSAY SCORE: {essay_total}"

    return {
        "grading_type": page_grades[0].get("grading_type", "STRICT_MATCH"),
        "objective_score_log": combined_objective_log,
        "essay_score_log": combined_essay_log,
        "confidence_score": avg_confidence,
        "objective_total": objective_total,
        "essay_total": essay_total,
        "score": objective_total + essay_total,
        "feedback": " | ".join(feedbacks),
    }

def transcribe_and_grade_page(img, page_num, context_block, cache_name=None):
    transcript_data = transcribe_single_page(img, page_num)
    transcribed_text = transcript_data.get("transcribed_text", "")
    grade_data = grade_single_page(transcribed_text, context_block, page_num, cache_name)
    grade_data["transcribed_text"] = transcribed_text
    return grade_data

#=========================================
#PARAREL RUN TRANSCRIPTION FOR TIER 1
#=========================================
def run_transcription(file_data):
    """
    Core transcription logic used by both sync and async routes.
    file_data: list of (filename, file_bytes) tuples.
    """
    all_images = []
    for filename, file_bytes in file_data:
        all_images.extend(read_file_as_images(file_bytes, label=filename or "exam"))

    print(f"⚙️ Transcribing {len(all_images)} page(s) in parallel...", flush=True)

    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        futures = [
            executor.submit(transcribe_single_page, img, i + 1)
            for i, img in enumerate(all_images)
        ]
        page_results = [f.result() for f in concurrent.futures.as_completed(futures)]

    data = merge_transcriptions(page_results)
    save_transcript_log(data)
    return data


# ==========================================
# JOB QUEUE (for async endpoints)
# ==========================================
job_store = {}  # job_id -> {"status": "queued|processing|done|error", "result": ..., "error": ...}
job_queue = Queue()


def worker():
    """Background thread that picks jobs off the queue and runs them."""
    while True:
        job_id, fn, args, kwargs = job_queue.get()
        job_store[job_id] = {"status": "processing"}
        try:
            result = fn(*args, **kwargs)
            job_store[job_id] = {"status": "done", "result": result}
        except Exception as e:
            print(f"❌ Job {job_id} failed: {e}")
            job_store[job_id] = {"status": "error", "error": str(e)}
        finally:
            job_queue.task_done()


NUM_WORKERS = 3
for _ in range(NUM_WORKERS):
    t = threading.Thread(target=worker, daemon=True)
    t.start()

print(f"🚀 Job queue started with {NUM_WORKERS} workers")


# ==========================================
# ROUTES
# ==========================================
@app.route('/ping', methods=['GET'])
def ping():
    print("🏓 PING received!")
    return jsonify({"status": "ok"}), 200


@app.route('/transcribe', methods=['POST'])
def transcribe():
    """Synchronous transcription — processes all pages in parallel, then returns."""
    files = request.files.getlist('file')
    if not files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    try:
        file_data = []
        for file in files:
            file_bytes = file.read()
            print(f"📄 File received: {file.filename}, size: {len(file_bytes)} bytes", flush=True)
            file_data.append((file.filename, file_bytes))

        data = run_transcription(file_data)
        return jsonify({"success": True, "data": data})

    except google.api_core.exceptions.ResourceExhausted:
        return jsonify({
            "success": False,
            "error": "quota_exceeded",
            "message": "AI service is temporarily unavailable. Please try again in a minute."
        }), 429
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        print(f"❌ Transcription Error: {e}", flush=True)
        return jsonify({"success": False, "error": "server_error", "message": "Something went wrong."}), 500


@app.route('/transcribe/async', methods=['POST'])
def transcribe_async():
    """Async transcription — enqueues the job and returns a job_id immediately."""
    files = request.files.getlist('file')
    if not files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    # Read bytes NOW before the request context closes
    file_data = [(f.filename, f.read()) for f in files]

    job_id = str(uuid.uuid4())
    job_store[job_id] = {"status": "queued"}
    job_queue.put((job_id, run_transcription, [file_data], {}))

    print(f"📥 Job {job_id} queued ({len(file_data)} file(s))")
    return jsonify({"success": True, "job_id": job_id}), 202


@app.route('/grade/async', methods=['POST'])
def grade_async():
    """Async version — enqueues transcribe+grade job, returns job_id immediately."""
    files = request.files.getlist('file')
    if not files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    # Read bytes NOW before request context closes
    file_data = [(f.filename, f.read()) for f in files]

    # Read all form fields now too
    context = request.form.get('context', '').strip()
    answer_key_text = request.form.get('answer_key_text', '').strip()
    reference_url = request.form.get('reference_url', '')
    reference_urls = request.form.getlist('reference_urls')
    answer_key_url = request.form.get('answer_key_url', '')
    answer_key_urls = request.form.getlist('answer_key_urls')

    def run_grade_job():
        all_answer_key_urls = list({answer_key_url, *answer_key_urls} - {''})
        all_reference_urls = list({reference_url, *reference_urls} - {''})

        context_block = ""
        extra_images = []

        if context:
            context_block += f"=== PROFESSOR CONTEXT ===\n{context}\n=========================\n\n"
        if answer_key_text:
            context_block += f"=== ESSAY RUBRIC ===\n{answer_key_text}\n===================\n\n"

        for url in all_answer_key_urls:
            try:
                imgs, text = fetch_and_parse_url(url)
                if text:
                    context_block += f"=== OBJECTIVE ANSWER KEY ===\n{text}\n============================\n\n"
                    print(f"📋 OBJECTIVE ANSWER KEY: {len(text)} chars appended")
                if imgs:
                  extracted_text = extract_images_to_text(imgs, label="Answer Key")
                  if extracted_text:
                      context_block += f"=== EXTRACTED REFERENCE IMAGES ===\n{extracted_text}\n==================================\n\n"
            except Exception as e:
                print(f"⚠️ Could not fetch answer key URL {url}: {e}")

        for url in all_reference_urls:
            try:
                imgs, text = fetch_and_parse_url(url)
                if text:
                    context_block += f"=== REFERENCE MATERIAL ===\n{text}\n==========================\n\n"
                    print(f"📋 REFERENCE MATERIAL: {len(text)} chars appended")
                if imgs:
                  extracted_text = extract_images_to_text(imgs, label="Answer Key")
                  if extracted_text:
                      context_block += f"=== EXTRACTED REFERENCE IMAGES ===\n{extracted_text}\n==================================\n\n"
            except Exception as e:
                print(f"⚠️ Could not fetch reference URL {url}: {e}")

        if not context_block.strip():
            context_block = "=== NO RUBRIC PROVIDED — Score everything as 0 ==="

        print(f"📋 Context block ready ({len(context_block)} chars)")

        # Build all images from all files
        all_images = []
        for filename, file_bytes in file_data:
            all_images.extend(read_file_as_images(file_bytes, label=filename or "exam"))

        print(f"⚙️ Transcribing + grading {len(all_images)} page(s) in parallel...")

       # After building context_block, before parallel grading:
        cache_name = create_context_cache(context_block)

        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
                futures = [
                    executor.submit(transcribe_and_grade_page, img, i + 1, context_block, cache_name)
                    for i, img in enumerate(all_images)
                ]
                page_grades = [f.result() for f in concurrent.futures.as_completed(futures)]
        finally:
            # Always clean up cache after grading, even if something errors
            if cache_name:
                delete_context_cache(cache_name)

        # Save merged transcript log
        merged_transcript = "\n\n".join(
            f"[PAGE {g['page']}]\n{g.get('transcribed_text', '')}"
            for g in sorted(page_grades, key=lambda x: x["page"])
        )
        
        avg_confidence = sum(g.get("confidence_score", 0) or 0 for g in page_grades) // len(page_grades)
        save_transcript_log({"transcribed_text": merged_transcript, "confidence_score": avg_confidence})

        data = merge_grades(page_grades)
        save_grade_log(data)

        print(f"✅ All pages graded. Final score={data['score']}")
        return data

    job_id = str(uuid.uuid4())
    job_store[job_id] = {"status": "queued"}
    job_queue.put((job_id, run_grade_job, [], {}))

    print(f"📥 Grade job {job_id} queued ({len(file_data)} file(s))")
    return jsonify({"success": True, "job_id": job_id}), 202


@app.route('/job/<job_id>', methods=['GET'])
def get_job(job_id):
    """Poll the status of an async job."""
    job = job_store.get(job_id)
    if not job:
        return jsonify({"success": False, "error": "Job not found"}), 404
    return jsonify({"success": True, **job})


@app.route('/grade', methods=['POST'])
def grade():
    try:
        # --- Support both JSON body (old) and multipart form with files (new) ---
        files = request.files.getlist('file')

        if files:
            # PATH B: files uploaded — transcribe + grade each page in parallel
            context = request.form.get('context', '').strip()
            answer_key_text = request.form.get('answer_key_text', '').strip()
            reference_url = request.form.get('reference_url', '')
            reference_urls = request.form.getlist('reference_urls')
            answer_key_url = request.form.get('answer_key_url', '')
            answer_key_urls = request.form.getlist('answer_key_urls')
            transcribed_text_override = ''
        else:
            # PATH A: JSON body with pre-transcribed text (original behavior)
            body = request.get_json(silent=True) or {}
            transcribed_text_override = body.get('transcribed_text', '').strip()
            context = body.get('context', '').strip()
            answer_key_text = body.get('answer_key_text', '').strip()
            reference_url = body.get('reference_url', '')
            reference_urls = body.get('reference_urls', [])
            answer_key_url = body.get('answer_key_url', '')
            answer_key_urls = body.get('answer_key_urls', [])

        all_answer_key_urls = list({answer_key_url, *answer_key_urls} - {''})
        all_reference_urls = list({reference_url, *reference_urls} - {''})

        # --- BUILD CONTEXT BLOCK (shared across all pages) ---
        context_block = ""
        extra_images = []

        if context:
            context_block += f"=== PROFESSOR CONTEXT ===\n{context}\n=========================\n\n"
        if answer_key_text:
            context_block += f"=== ESSAY RUBRIC ===\n{answer_key_text}\n===================\n\n"

        for url in all_answer_key_urls:
            try:
                imgs, text = fetch_and_parse_url(url)
                if text:
                    context_block += f"=== OBJECTIVE ANSWER KEY ===\n{text}\n============================\n\n"
                    print(f"📋 OBJECTIVE ANSWER KEY: {len(text)} chars appended")
                if imgs:
                  extracted_text = extract_images_to_text(imgs, label="Answer Key")
                  if extracted_text:
                      context_block += f"=== EXTRACTED REFERENCE IMAGES ===\n{extracted_text}\n==================================\n\n"
            except Exception as e:
                print(f"⚠️ Could not fetch answer key URL {url}: {e}")

        for url in all_reference_urls:
            try:
                imgs, text = fetch_and_parse_url(url)
                if text:
                    context_block += f"=== REFERENCE MATERIAL ===\n{text}\n==========================\n\n"
                    print(f"📋 REFERENCE MATERIAL: {len(text)} chars appended")
                if imgs:
                  extracted_text = extract_images_to_text(imgs, label="Answer Key")
                  if extracted_text:
                      context_block += f"=== EXTRACTED REFERENCE IMAGES ===\n{extracted_text}\n==================================\n\n"
            except Exception as e:
                print(f"⚠️ Could not fetch reference URL {url}: {e}")

        if not context_block.strip():
            context_block = "=== NO RUBRIC PROVIDED — Score everything as 0 ==="

        print(f"📋 Final context_block ({len(context_block)} chars):\n{context_block[:600]}")

        # --- PATH A: pre-transcribed text, grade directly ---
        if transcribed_text_override:
            if not transcribed_text_override:
                return jsonify({"success": False, "error": "No transcription provided"}), 400

            data = grade_single_page(transcribed_text_override, context_block, page_num=1)
            save_grade_log(data)

            obj = data.get("objective_total", 0) or 0
            essay_from_json = data.get("essay_total", 0) or 0
            essay_from_log = extract_essay_total_from_log(data.get("essay_score_log", ""))
            essay = max(essay_from_json, essay_from_log)

            if essay_from_json != essay_from_log:
                print(f"⚠️ Essay total mismatch (PATH A) — JSON: {essay_from_json}, Log: {essay_from_log}")

            data["essay_total"] = essay
            data["score"] = obj + essay

            print(f"✅ Grading done: score={data.get('score')}")
            return jsonify({"success": True, "data": data})

        # --- PATH B: files uploaded — transcribe + grade each page in parallel ---
        if not files:
            return jsonify({"success": False, "error": "No file or transcribed_text provided"}), 400

        all_images = []
        for f in files:
            file_bytes = f.read()
            print(f"📄 File received: {f.filename}, size: {len(file_bytes)} bytes")
            all_images.extend(read_file_as_images(file_bytes, label=f.filename or "exam"))

        print(f"⚙️ Transcribing + grading {len(all_images)} page(s) in parallel...")

        # After building context_block, before parallel grading:
        cache_name = create_context_cache(context_block)

        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
                futures = [
                    executor.submit(transcribe_and_grade_page, img, i + 1, context_block, cache_name)
                    for i, img in enumerate(all_images)
                ]
                page_grades = [f.result() for f in concurrent.futures.as_completed(futures)]
        finally:
            # Always clean up cache after grading, even if something errors
            if cache_name:
                delete_context_cache(cache_name)

        # Save merged transcript log
        merged_transcript = "\n\n".join(
            f"[PAGE {g['page']}]\n{g.get('transcribed_text', '')}"
            for g in sorted(page_grades, key=lambda x: x["page"])
        )
        avg_confidence = sum(g.get("confidence_score", 0) or 0 for g in page_grades) // len(page_grades)
        save_transcript_log({"transcribed_text": merged_transcript, "confidence_score": avg_confidence})

        data = merge_grades(page_grades)
        save_grade_log(data)

        print(f"✅ All pages graded. Final score={data['score']}")
        return jsonify({"success": True, "data": data})

    except google.api_core.exceptions.ResourceExhausted:
        return jsonify({
            "success": False,
            "error": "quota_exceeded",
            "message": "AI service is temporarily unavailable. Please try again in a minute."
        }), 429
    except Exception as e:
        print(f"❌ Grading Error: {e}")
        return jsonify({"success": False, "error": "server_error", "message": "Something went wrong."}), 500


@app.route('/masterlist', methods=['POST'])
def masterlist():
    """Extract student names and IDs from an uploaded image or PDF."""
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    try:
        import tempfile
        file = request.files['file']
        file_bytes = file.read()
        filename = file.filename or "masterlist"
        print(f"📄 Masterlist file: {filename}, {len(file_bytes)} bytes")

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        uploaded_file = genai.upload_file(tmp_path, mime_type="application/pdf")
        os.unlink(tmp_path)

        ai_inputs = [
            uploaded_file,
            """Extract ALL student data from this PDF as JSON.
Return ONLY a JSON array, no extra text:
[{"name": "FullName", "id": "StudentID_or_null"}]
Include every student listed. Student IDs are typically in format like TUPM-XX-XXXX."""
        ]

        response = transcription_model.generate_content(
            ai_inputs,
            generation_config=GENERATION_CONFIG,
            safety_settings=SAFETY_SETTINGS,
        )

        raw_text = response.text
        print(f"📋 AI response: {raw_text[:300]}")
        start = raw_text.find('[')
        end = raw_text.rfind(']') + 1

        if start != -1:
            data = json.loads(raw_text[start:end])
            print(f"👥 Extracted {len(data)} students")
            return jsonify({"success": True, "data": data})
        else:
            return jsonify({"success": False, "error": "Could not parse masterlist"}), 500

    except google.api_core.exceptions.ResourceExhausted:
        return jsonify({
            "success": False,
            "error": "quota_exceeded",
            "message": "AI service is temporarily unavailable. Please try again in a minute."
        }), 429
    except Exception as e:
        print(f"❌ Masterlist Error: {e}")
        return jsonify({"success": False, "error": "server_error", "message": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)